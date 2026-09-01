from pathlib import Path

from pypdf import PdfReader
from docx import Document


class DocumentLoader:

    @staticmethod
    def load(filepath: str) -> str:

        extension = Path(filepath).suffix.lower()

        loaders = {
            ".pdf": DocumentLoader._load_pdf,
            ".docx": DocumentLoader._load_docx,
            ".txt": DocumentLoader._load_txt,
            ".md": DocumentLoader._load_txt,
            ".csv": DocumentLoader._load_txt,
        }

        loader = loaders.get(extension)

        if loader is None:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported formats: PDF, DOCX, TXT, MD, CSV."
            )

        text = loader(filepath)

        if not text or not text.strip():
            raise ValueError(
                "No readable text could be extracted from this document."
            )

        return text.strip()

    # ---------------------------------
    # PDF
    # ---------------------------------

    @staticmethod
    def _load_pdf(filepath: str) -> str:

        reader = PdfReader(filepath)

        pages = []

        for page_number, page in enumerate(reader.pages, start=1):

            extracted = page.extract_text()

            if extracted:
                pages.append(
                    f"[Page {page_number}]\n{extracted.strip()}"
                )

        return "\n\n".join(pages)

    # ---------------------------------
    # DOCX
    # ---------------------------------

    @staticmethod
    def _load_docx(filepath: str) -> str:

        doc = Document(filepath)

        sections = []

        # Paragraphs
        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                sections.append(text)

        # Tables
        for table_index, table in enumerate(
            doc.tables,
            start=1
        ):

            rows = []

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                rows.append(
                    " | ".join(cells)
                )

            table_text = "\n".join(
                row for row in rows if row.strip()
            )

            if table_text:
                sections.append(
                    f"[Table {table_index}]\n"
                    f"{table_text}"
                )

        return "\n\n".join(sections)

    # ---------------------------------
    # TXT / MD / CSV
    # ---------------------------------

    @staticmethod
    def _load_txt(filepath: str) -> str:

        with open(
            filepath,
            "r",
            encoding="utf-8",
            errors="replace",
        ) as file:

            return file.read()